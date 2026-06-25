"""
Generate a comprehensive methods-grade Word document for ABEL.
Run with:
    "c:/Users/jober/Desktop/ABEL realism/.venv/Scripts/python.exe" generate_readme_docx.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Heading styles ────────────────────────────────────────────────────────────
styles = doc.styles

h1 = styles['Heading 1']
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

h2 = styles['Heading 2']
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

h3 = styles['Heading 3']
h3.font.size = Pt(11)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

normal = styles['Normal']
normal.font.size = Pt(11)
normal.font.name = 'Calibri'

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.style = doc.styles['Normal']
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.left_indent = Inches(0.5)
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = str(val)
    return row

def section_break():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("ABEL")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_para.add_run(
    "Active Learning-Driven Behavioral Scoring from Pose-Tracking Data\n"
    "Technical Reference and Methods Documentation"
)
run2.font.size = Pt(14)
run2.italic = True

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.add_run(f"Document generated: {datetime.date.today().strftime('%B %d, %Y')}")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Contents", 1)
toc_items = [
    ("1", "Overview and Design Philosophy"),
    ("2", "System Architecture"),
    ("3", "Project Structure and Data Organisation"),
    ("4", "Data Import and Session Linking"),
    ("5", "Pose Data Cleaning and Smoothing"),
    ("6", "Extracting Features Tab: Pose Feature Extraction"),
    ("7", "Extracting Features Tab: Context Feature Extraction"),
    ("8", "Motif Discovery: Unsupervised Kinematic Clustering"),
    ("9", "Candidate Generation"),
    ("10", "Clip Extraction and Preprocessing"),
    ("11", "Seed Examples and Behavior Signature Building"),
    ("12", "Behavior Representation: Segment-Level Feature Vectors"),
    ("13", "Active Learning Tab: Closed-Loop Modeling Pipeline"),
    ("    13.1", "Pipeline Stages"),
    ("    13.2", "Classifier Training"),
    ("    13.3", "Probability Calibration"),
    ("    13.4", "Uncertainty Scoring"),
    ("    13.5", "Context Feature Integration in Active Learning"),
    ("    13.6", "Video-Based Fusion Inference"),
    ("    13.7", "Candidate Ranking Strategies"),
    ("14", "Review Tab: Human-in-the-Loop Labeling"),
    ("15", "Temporal Refinement: Dense Sliding-Window Inference"),
    ("16", "Temporal Review Tab: Bout Extraction and Threshold Tuning"),
    ("17", "Model Evaluation Metrics"),
    ("18", "Phase 1 Adaptive Benchmarking (Advanced)"),
    ("19", "Export"),
    ("20", "Provenance and Reproducibility"),
    ("21", "Software Dependencies"),
    ("22", "Glossary"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.add_run(f"{num}  {title}")
    p.paragraph_format.left_indent = Inches(0.1 if not num.startswith(" ") else 0.4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Overview and Design Philosophy", 1)

add_para(
    "ABEL is a desktop application for semi-automated behavioral scoring of video recordings "
    "generated by standard tracking systems (e.g., DeepLabCut). It combines unsupervised kinematic "
    "analysis, supervised machine learning, and iterative human review into a single closed-loop "
    "workflow. The core goal is to reduce the manual annotation burden in behavioral neuroscience "
    "by progressively refining a probabilistic behavior classifier through active learning — "
    "directing the researcher's labeling effort toward the segments where the model is most "
    "uncertain or where additional labels would most improve its performance."
)

add_para(
    "The application is behavior-agnostic: it is designed to detect any temporally discrete "
    "behavior of interest defined by the researcher, given an operational definition, a small "
    "set of seed examples, and iterative labeled feedback. ABEL does not require manual "
    "frame-by-frame annotation; instead, it presents short candidate video clips ranked by "
    "predicted relevance and model uncertainty for accept/reject decisions."
)

add_heading("Key Design Principles", 2)
add_bullet("No-video-first feature extraction: kinematic features are derived entirely from pose tracking files "
           "(no video decoding) until the optional clip extraction and context feature steps.")
add_bullet("Crash-resilient persistence: all project state, model artifacts, and training snapshots are "
           "stored atomically on disk so work is not lost on unexpected exits.")
add_bullet("Graceful degradation: every optional dependency (GPU, CUDA, UMAP, LightGBM, etc.) has a "
           "CPU fallback so the pipeline can run on any machine.")
add_bullet("Modular service layer: UI tabs communicate with independent services; all computation "
           "logic is outside the UI for testability.")
add_bullet("Provenance tracking: every derived artifact records the config hash, model version, "
           "feature version, and timestamp that produced it.")

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. System Architecture", 1)

add_para(
    "ABEL is structured in six horizontal layers:"
)

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = "Layer"
hdr[1].text = "Purpose"
for c in hdr:
    for run in c.paragraphs[0].runs:
        run.bold = True

layer_rows = [
    ("UI (abel/ui)", "Qt6 (PySide6) widgets, tabs, and dialogs. Each tab owns visual state "
     "only; computation is delegated to services."),
    ("Services (abel/services)", "All domain logic: feature extraction, model training, "
     "candidate ranking, evaluation, export, etc. Services are stateless except for a project_root path."),
    ("Workers (abel/workers)", "Background QRunnable wrappers that move service calls off "
     "the Qt main thread to keep the UI responsive."),
    ("Storage (abel/storage)", "Atomic file I/O helpers for YAML, JSON, Parquet, and NumPy "
     "artifacts. Write-to-temp-then-rename pattern prevents partial writes."),
    ("Models (abel/models)", "Pydantic v2 data models for all key entities: sessions, "
     "behaviors, candidates, labels, model cards, presets, etc."),
    ("Core (abel/core)", "App-wide constants and custom exception types."),
    ("Temporal Refinement (abel/temporal_refinement)", "Dense per-frame inference pipeline "
     "and bout post-processing, separate from the active-learning loop."),
]
for row_data in layer_rows:
    add_table_row(tbl, row_data)

section_break()
add_para(
    "The overall data flow through the pipeline is as follows:"
)
pipeline_steps = [
    "Data Import → link video + pose files into sessions",
    "Behavior Definitions → define target behaviors with operational criteria",
    "Seed Examples → annotate 2–10 confirmed positive examples per behavior",
    "Pose Feature Extraction → compute kinematic windows from pose CSV/H5 files",
    "Motif Discovery → unsupervised clustering of kinematic windows",
    "Candidate Generation → rank motif windows for clip extraction",
    "Clip Extraction → decode video only for top-ranked windows",
    "Active Learning Tab → build representation → train model → score uncertainty → rank candidates",
    "Review Tab → present clips to researcher; researcher labels accept/reject/relabel",
    "[ repeat Active Learning + Review until F1 / PR-AUC converge ]",
    "Temporal Refinement → dense sliding-window inference across full recordings",
    "Temporal Review → threshold tuning, bout extraction, false-positive flagging",
    "Export → CSV / XLSX per-subject bout tables",
]
for step in pipeline_steps:
    add_numbered(step)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Project Structure and Data Organisation", 1)

add_para(
    "Each ABEL project is a directory tree. On creation, the following subdirectories are "
    "created automatically:"
)

dir_tbl = doc.add_table(rows=1, cols=2)
dir_tbl.style = 'Table Grid'
dir_tbl.rows[0].cells[0].text = "Path (project-relative)"
dir_tbl.rows[0].cells[1].text = "Contents"
for c in dir_tbl.rows[0].cells:
    for run in c.paragraphs[0].runs:
        run.bold = True

dir_rows = [
    ("config/", "YAML configuration files: behavior definitions, preprocessing presets, "
     "model settings, ROI definitions, export settings."),
    ("raw/videos/, raw/pose/", "Source video and pose files (symlinked or copied depending on import mode)."),
    ("derived/pose_features/", "Per-session compressed NumPy .npz files: kinematic window feature matrices."),
    ("derived/context_features/", "Per-session Parquet files: optical-flow and spatial context "
     "features at frame resolution."),
    ("derived/representations/", "Merged segment-level feature matrices (frame_features.parquet, "
     "segment_features.parquet) used as classifier input."),
    ("derived/models/", "Trained model artifacts: model_state.pkl, model_card.yaml, "
     "validation_predictions.parquet."),
    ("derived/training_sets/", "Accumulated training set Parquet file and versioned snapshots."),
    ("derived/review_labels/", "Reviewer label records (reviewer_labels.parquet)."),
    ("derived/clips/", "Extracted video clips for review."),
    ("derived/behavior_bouts/", "Merged bout tables per behavior (Parquet)."),
    ("derived/evaluation/", "Metrics JSON, PR curve PNG, confusion matrix PNG, UMAP/PCA separation plot."),
    ("derived/temporal_refinement/", "Per-behavior per-session dense probability traces and bout outputs."),
    ("derived/analysis/benchmarks/, derived/analysis/diagnostics/", "Phase 1 adaptive benchmarking "
     "outputs: CSV tables, diagnostic PNG/SVG plots."),
    ("exports/", "Final CSV, XLSX, and Parquet exports for downstream analysis."),
]
for row_data in dir_rows:
    add_table_row(dir_tbl, row_data)

add_para(
    "Project configuration is split into a ProjectConfig (YAML) for user-editable experiment "
    "parameters and a ProjectState (JSON) for runtime state such as the active tab and review "
    "progress. Both use Pydantic v2 models with schema versioning for forward compatibility."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. DATA IMPORT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. Data Import and Session Linking", 1)

add_para(
    "The Data Import tab ingests video files and their corresponding pose-tracking CSV or H5 "
    "files exported from DeepLabCut (or compatible formats). The importer automatically pairs "
    "video and pose files by fuzzy filename matching, producing linked sessions stored in "
    "derived/review_tables/import_manifest.json."
)

add_heading("Session and Subject Identification", 2)
add_para(
    "Subject identifiers are extracted from filenames using a configurable regular expression "
    "(default: the leading alphabetic+numeric prefix, e.g., 'TMT2' from 'TMT2DLC_...'). "
    "Each unique video+pose pairing becomes one session record with a stable session_id. "
    "Pixels-per-millimetre (px/mm) can be entered per session to convert pixel-unit kinematic "
    "features to physical units (mm/s for speed, mm for displacement)."
)

add_heading("Supported Pose Formats", 2)
add_bullet("DeepLabCut CSV export: header rows encoding scorer name, body-part names, and x/y/likelihood "
           "coordinates for each frame.")
add_bullet("DeepLabCut H5 export: equivalent hierarchical DataFrame structure.")
add_para(
    "The PoseProcessingService normalises both formats into an internal PoseData representation "
    "containing per-body-part x, y, and likelihood arrays indexed by frame number."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. POSE DATA CLEANING
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Pose Data Cleaning and Smoothing", 1)

add_para(
    "Before any kinematic features are computed, each pose file is cleaned in the following steps. "
    "These parameters are configurable via the PoseSmoothingSettings in project config."
)

add_heading("Step 1: Likelihood Thresholding", 2)
add_para(
    "DeepLabCut assigns a per-frame likelihood score (0–1) to every keypoint detection. "
    "Detections with likelihood below the configured threshold (default: 0.2) are treated as "
    "missing — their x/y coordinates are set to NaN. This removes low-confidence detections "
    "caused by occlusion, model errors, or the animal leaving the frame."
)

add_heading("Step 2: Linear Interpolation of Dropout Gaps", 2)
add_para(
    "Consecutive NaN frames (caused by thresholding or tracking loss) shorter than a configurable "
    "maximum gap (default: 10 frames) are filled by linear interpolation between the nearest "
    "valid detections on either side. Gaps exceeding this length remain as NaN to flag "
    "extended tracking failures."
)

add_heading("Step 3: Centred Rolling-Average Smoothing", 2)
add_para(
    "A centred rolling mean with a configurable odd-numbered window (default: 5 frames, "
    "corresponding to ~167 ms at 30 fps) is applied to all x and y coordinate series. "
    "This attenuates high-frequency jitter from sub-pixel tracking noise while preserving "
    "genuine behaviorally relevant motion. A window of 1 disables smoothing."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. POSE FEATURE EXTRACTION TAB
# ══════════════════════════════════════════════════════════════════════════════
add_heading("6. Extracting Features Tab: Pose Feature Extraction", 1)

add_para(
    "The Pose Features tab (and its underlying PoseFeaturesService) computes kinematic feature "
    "vectors from the cleaned pose data across sliding windows of the entire recording, "
    "without decoding any video. The outputs are compressed NumPy .npz files stored in "
    "derived/pose_features/, one per session."
)

add_heading("Sliding Window Parameterisation", 2)
add_para(
    "The recording is segmented into overlapping windows defined by two parameters:"
)
add_bullet("Window duration (default: 2.0 s at 30 fps = 60 frames): the temporal extent of each feature window.")
add_bullet("Window stride (default: 1.0 s = 30 frames): the step size between successive window start frames.")

add_para(
    "Built-in presets include: Standard (2 s / 1 s stride), Short Window (1 s / 0.5 s), "
    "Long Window (4 s / 2 s), and High-Res (1 s / 0.25 s). Custom presets can be saved to "
    "config/pose_features.yaml."
)

add_heading("Feature Vector Layout", 2)
add_para(
    "For each window, a nine-dimensional feature vector is computed from the body centroid "
    "trajectory and body-axis orientation:"
)

feat_tbl = doc.add_table(rows=1, cols=3)
feat_tbl.style = 'Table Grid'
feat_tbl.rows[0].cells[0].text = "Feature"
feat_tbl.rows[0].cells[1].text = "Description"
feat_tbl.rows[0].cells[2].text = "Units"
for c in feat_tbl.rows[0].cells:
    for run in c.paragraphs[0].runs:
        run.bold = True

feat_rows = [
    ("speed_mean", "Mean instantaneous body centroid speed across the window", "mm/s or px/s"),
    ("speed_std", "Standard deviation of instantaneous speed", "mm/s or px/s"),
    ("speed_max", "Peak instantaneous speed across the window", "mm/s or px/s"),
    ("disp_mean", "Mean per-frame Euclidean displacement of the centroid", "mm or px"),
    ("disp_std", "Standard deviation of per-frame displacement", "mm or px"),
    ("axis_cos_mean", "Mean cosine of the body-axis heading angle — rotation-invariant orientation encoding", "dimensionless"),
    ("axis_sin_mean", "Mean sine of the body-axis heading angle", "dimensionless"),
    ("axis_angle_std", "Standard deviation of body-axis heading angle — captures postural variability", "radians"),
    ("likelihood_mean", "Mean DLC keypoint likelihood score across the window — tracking quality proxy", "dimensionless [0–1]"),
]
for row_data in feat_rows:
    add_table_row(feat_tbl, row_data)

add_para(
    "The body centroid is computed as the mean of all tracked keypoints with likelihood above "
    "the cleaning threshold. The body-axis angle is derived from the first principal axis of the "
    "keypoint cloud (proximal–distal axis). Speed is computed as the frame-to-frame Euclidean "
    "distance of the centroid, multiplied by the video frame rate to obtain physical velocity."
)

add_para(
    "When a px/mm calibration is available for the session (entered in Data Import), all "
    "distance and speed features are converted to physical units (mm, mm/s). Without "
    "calibration, features remain in pixel units and a warning is logged."
)

add_para(
    "Output: derived/pose_features/{session_id}.npz — a compressed archive containing the "
    "feature matrix (n_windows × 9), window frame start/end indices, feature names, and "
    "body-part names. A session summary JSON is also written for the UI."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONTEXT FEATURE EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("7. Extracting Features Tab: Context Feature Extraction", 1)

add_para(
    "The Context Feature tab runs the ContextFeatureService, which opens the raw video to "
    "compute per-frame environment-interaction features that complement the pose kinematics. "
    "These features capture local substrate motion at the animal's body position, paw-region "
    "and nose-region dynamics, and spatial proximity to user-defined regions of interest (ROIs). "
    "They are computed at frame resolution and stored as Parquet files in "
    "derived/context_features/sessions/{session_id}.parquet."
)

add_heading("Optical Flow (Farnebäck Dense Optical Flow)", 2)
add_para(
    "Dense optical flow is computed using OpenCV's Farnebäck algorithm between consecutive "
    "downsampled grayscale frames. The algorithm parameters are configurable:"
)
add_bullet("Pyramid scale: 0.5 (each level is half the resolution of the previous)")
add_bullet("Pyramid levels: 3")
add_bullet("Window size: 15 px")
add_bullet("Iterations: 3")
add_bullet("Polynomial expansion neighbourhood (poly_n): 5")
add_bullet("Polynomial sigma: 1.2")

add_para(
    "Spatial downsampling is applied automatically to target a ~512-pixel long edge, using the "
    "largest power-of-two factor that keeps the long edge at or above 512 pixels. This reduces "
    "computational cost by 4–16× on typical  high-definition videos while preserving "
    "motion information at behaviorally relevant spatial scales."
)

add_para(
    "When GPU-enabled OpenCV is available (CUDA build), the optical flow computation is "
    "offloaded to the GPU. A threading lock serialises GPU calls across parallel worker threads "
    "to prevent CUDA contention."
)

add_heading("Per-Frame Features Computed", 2)

ctx_tbl = doc.add_table(rows=1, cols=2)
ctx_tbl.style = 'Table Grid'
ctx_tbl.rows[0].cells[0].text = "Feature"
ctx_tbl.rows[0].cells[1].text = "Description"
for c in ctx_tbl.rows[0].cells:
    for run in c.paragraphs[0].runs:
        run.bold = True

ctx_rows = [
    ("local_surface_energy", "Mean foreground-mask intensity at the body centroid crop "
     "(radius ~12 px after downsampling) from a MOG2 background subtractor — proxies substrate "
     "disturbance at the animal's body position."),
    ("local_surface_var", "Pixel-intensity variance in the body centroid crop — captures texture "
     "complexity of the local substrate."),
    ("local_surface_change", "Frame-to-frame absolute pixel difference at the body centroid crop — "
     "quantifies local substrate motion."),
    ("nose_local_change", "Frame-to-frame absolute pixel difference in a tight crop around the "
     "nose keypoint — sensitive to snout contact/interaction with objects."),
    ("nose_local_variance", "Pixel-intensity variance in the nose crop."),
    ("flow_mag_paw_l / flow_mag_paw_r", "Mean optical-flow magnitude at the left and right forepaw "
     "keypoint crops — directly measures paw movement."),
    ("flow_mag_nose", "Mean optical-flow magnitude at the nose keypoint crop."),
    ("flow_mag_tmt / flow_mag_target", "Mean optical-flow magnitude inside the target ROI (e.g., "
     "stimulus object area) — captures object-directed movement."),
    ("flow_dir_paw", "Mean flow direction at the paw region relative to the body centroid — "
     "distinguishes forward vs. backward paw movement."),
    ("flow_entropy_local", "Shannon entropy of the optical-flow magnitude histogram at the body "
     "centroid crop — high entropy = complex, multi-directional local motion."),
    ("nose_to_target_dist", "Euclidean distance from the nose keypoint to the centre of the target ROI, "
     "in mm when px/mm is calibrated."),
    ("body_centroid_to_target_dist", "Euclidean distance from the body centroid to the target ROI centre."),
    ("paw_l/r_to_target_dist", "Euclidean distance from each forepaw keypoint to the target ROI centre."),
    ("angle_to_target", "Heading angle of the nose relative to the target ROI — captures approach orientation."),
]
for row_data in ctx_rows:
    add_table_row(ctx_tbl, row_data)

add_heading("Background Modelling (MOG2)", 2)
add_para(
    "A Mixture of Gaussians (MOG2) background subtractor is maintained per video chunk. "
    "To ensure the model is well-initialised, each processing chunk replays up to 50 preceding "
    "frames as a warm-up pass before recording feature values. The foreground mask produced by "
    "MOG2 at the animal's body position provides a substrate-disturbance signal — e.g., freshly "
    "disturbed bedding under the animal — that is independent of pose tracking quality."
)

add_heading("Keypoint Role Mapping", 2)
add_para(
    "Context features are conditioned on specific anatomical roles (nose, left forepaw, right "
    "forepaw). The service resolves these roles from the actual DLC keypoint names by tokenising "
    "each body-part name (splitting on underscores, hyphens, and spaces) and matching against "
    "a priority-ordered vocabulary (e.g., 'forepaw_left' → paw_l role). This makes the mapping "
    "fully data-driven and compatible with any DLC labeling convention without hardcoding exact "
    "keypoint names."
)

add_heading("Parallel Processing", 2)
add_para(
    "Long recordings are split into chunks of ~2000 frames and processed in parallel using a "
    "ThreadPoolExecutor. Each chunk opens its own VideoCapture object. Optical-flow continuity "
    "across chunk boundaries is preserved by reading the first frame before each chunk start "
    "as the initial prev_gray. Because OpenCV's Farnebäck implementation releases the Python "
    "GIL, threads genuinely overlap CPU work."
)
section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. MOTIF DISCOVERY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("8. Motif Discovery: Unsupervised Kinematic Clustering", 1)

add_para(
    "The Motif Discovery step clusters pose-feature windows into kinematic 'motifs' — recurring "
    "movement patterns — using unsupervised learning. This step does not require any behavioral "
    "labels; it is analogous to a data-driven ethogram prior. No video is decoded."
)

add_heading("Algorithms", 2)

add_bullet("K-Means (default): scikit-learn MiniBatchKMeans on standardised feature vectors. "
           "Default: 10 or 20 clusters.")
add_bullet("UMAP + K-Means: UMAP dimensionality reduction (configurable n_components, n_neighbors, "
           "min_dist) followed by K-Means on the low-dimensional coordinates.")
add_bullet("UMAP + HDBSCAN: UMAP reduction followed by density-based HDBSCAN clustering (auto "
           "number of clusters). Noise-labelled windows (HDBSCAN cluster = -1) are saved as a "
           "'noise' motif and excluded from candidate generation by default.")

add_heading("Feature Standardisation", 2)
add_para(
    "Before clustering, all feature dimensions are z-scored (mean-subtracted, divided by "
    "standard deviation) across the combined window set. When seed examples are provided, "
    "standardisation is fit on seed-overlapping windows only and then applied to all windows. "
    "This biases the feature space toward the kinematic regime of the target behavior while "
    "still assigning motif labels to the full recording."
)

add_heading("Seed-Focused Clustering", 2)
add_para(
    "When seed examples are supplied, the clustering model is trained exclusively on pose "
    "windows that temporally overlap with the seed annotations (overlap determined by "
    "frame-interval intersection). The trained model then assigns motif labels to every "
    "window in the full recording via predict(). This focused training prevents the cluster "
    "centres from being dominated by common non-target movement patterns and instead "
    "resolves the kinematic sub-space near the behavior of interest."
)

add_heading("Seed-Motif Prior", 2)
add_para(
    "After clustering, each motif is assigned a seed-motif prior score: the fraction of "
    "windows belonging to that motif that overlap with any seed example. Motifs with high "
    "prior scores are enriched for the target behavior and will receive higher weight during "
    "candidate generation."
)

add_heading("Outputs", 2)
add_para(
    "Outputs are stored in derived/motifs/: a motif model YAML (cluster centres, algorithm "
    "parameters), a motif assignments JSON (session_id, start_frame, end_frame, motif_id, "
    "confidence for every window), and a cluster summary with per-motif counts and seed "
    "overlap rates."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. CANDIDATE GENERATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("9. Candidate Generation", 1)

add_para(
    "Candidate Generation ranks motif-assigned windows to identify the most promising "
    "time segments for video clip extraction and review. It operates on the motif assignment "
    "file and seed examples produced by the previous steps."
)

add_heading("Episode Construction", 2)
add_para(
    "Individual sliding windows from pose feature extraction are merged into 'episodes': "
    "variable-length temporal segments composed of consecutive windows assigned to the same "
    "motif. Consecutive same-motif windows within a session-specific merge gap are fused into "
    "one episode. This prevents the user from seeing many nearly-identical short clips "
    "corresponding to a single continuous behavioral event."
)

add_heading("Composite Scoring", 2)
add_para(
    "Each episode is assigned a composite score from two components:"
)
add_bullet("Motif score (default weight: 0.7): combines the motif cluster assignment confidence "
           "and the seed-motif prior (fraction of windows in that motif that overlapped seeds). "
           "Formula: motif_score = 0.7 × cluster_confidence + 0.3 × motif_seed_prior.")
add_bullet("Seed similarity score (default weight: 0.3): the maximum temporal overlap fraction "
           "between the candidate episode and any seed example (Jaccard-style frame overlap). "
           "Windows that directly co-occur with seeds receive the highest seed scores.")
add_para(
    "Total score = 0.7 × motif_score + 0.3 × seed_similarity_score."
)

add_heading("Non-Overlapping Selection", 2)
add_para(
    "After scoring, candidates are sorted in descending order and selected greedily to avoid "
    "temporal redundancy: a candidate is accepted only if its overlap with any already-selected "
    "candidate (computed as the intersection length divided by the shorter episode length) is "
    "below a configurable threshold (default: 0.6). The top-k non-overlapping candidates "
    "(default: 300 per run) are retained."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. CLIP EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("10. Clip Extraction and Preprocessing", 1)

add_para(
    "The Preprocessing tab triggers the ClipExtractionService, which decodes video only for "
    "the candidate windows selected by the previous step. This selective approach avoids "
    "processing the entire recording (which may be hours long) and instead targets only the "
    "brief segments most likely to contain the behavior of interest."
)

add_heading("Clip Output Presets", 2)

add_para(
    "Three built-in presets define the clip output format; custom presets can be saved:"
)
add_bullet("Standard Crop (224×224, 15 fps)")
add_bullet("Large Crop (256×256, 15 fps)")
add_bullet("Grayscale (128×128, 10 fps)")

add_heading("Animal-Centred Cropping", 2)
add_para(
    "For each frame within a clip, the body centroid (estimated from the pose keypoints "
    "after cleaning) is used as the crop centre. A configurable margin (default: 80 px) "
    "is added around the centroid on all sides. The crop margin scales with the source video "
    "resolution: higher-resolution inputs receive proportionally larger crops so that the "
    "animal occupies a consistent fraction of the output frame regardless of recording setup."
)

add_para(
    "When adaptive cropping is enabled (default: on), the crop margin is additionally scaled "
    "by the square root of the crop_area_scale factor (default: 1.25) to control the "
    "proportion of the scene visible around the animal."
)

add_heading("Optional Post-Processing", 2)
add_bullet("Body-axis rotation: rotates each frame so the animal's head-to-tail axis is consistently "
           "oriented, removing viewing-angle variance before review or further learning.")
add_bullet("Stabilization: translates frames to keep the body centroid centred in the output.")
add_bullet("Grayscale conversion: reduces clip file size for large datasets.")

add_para(
    "Clips are written to derived/clips/ with deterministic, collision-safe filenames "
    "(ASCII-safe stem + 8-character SHA-1 digest of the clip ID). A ClipManifest JSON "
    "records the clip ID, session ID, score components, and preprocessing settings for "
    "provenance tracking."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. SEED EXAMPLES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("11. Seed Examples and Behavior Signature Building", 1)

add_para(
    "Seed examples are manually annotated positive instances of the target behavior. They serve "
    "as the initial supervisory signal before any labeled training data exists."
)

add_heading("Seed Annotation", 2)
add_para(
    "In the Seed Examples tab, the researcher plays through an imported session and marks the "
    "start and end frame of 2–10 unambiguous examples of the target behavior. Each seed record "
    "stores: session_id, start_frame, end_frame, label_type (positive/negative), quality_flag, "
    "and optional notes."
)

add_heading("Behavior Signature Building", 2)
add_para(
    "The BehaviorSignatureBuilder uses seed annotations and motif assignments to derive a "
    "BehaviorSignature object describing the kinematic profile of the target behavior:"
)
add_bullet("Enriched syllables: motif/syllable IDs overrepresented in seed-overlapping windows, "
           "with frequency-based scores.")
add_bullet("Depleted syllables: motif IDs expected from the overall distribution but absent from seeds.")
add_bullet("Syllable sequences: the most common consecutive motif bigrams and trigrams within seed windows.")
add_bullet("Transition matrix: syllable-to-syllable transition probability matrix from seed intervals.")
add_bullet("Duration statistics: mean, median, SD, min, max behavior bout duration in seconds "
           "estimated from seed intervals.")
add_bullet("Pose constraints: mean kinematic feature values (speed, displacement, body-axis "
           "variability, tracking quality) computed over seed-overlapping pose-feature windows.")

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. BEHAVIOR REPRESENTATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("12. Behavior Representation: Segment-Level Feature Vectors", 1)

add_para(
    "Before classifier training, per-frame pose and context features are aggregated into "
    "segment-level feature vectors by the BehaviorRepresentationService. These vectors "
    "form the input matrix for the active-learning classifier."
)

add_heading("Frame-Level Data Merging", 2)
add_para(
    "Per-frame pose features (from derived/context_features/) and per-frame context features "
    "are joined on the (session_id, frame) key. Missing context features are filled with "
    "session-mean values to handle recordings where context features have not been computed."
)

add_heading("Per-Group Z-Score Normalisation", 2)
add_para(
    "Feature columns are z-scored independently within each (animal_id, session_id) group. "
    "The group-wise standardisation removes systematic differences in scale across recording "
    "sessions (e.g., different arenas, lighting conditions) and ensures that the model "
    "generalises to unseen sessions. Columns with zero variance or single-sample groups "
    "receive a sigma of 1.0 to avoid division by zero."
)

add_heading("Segment Window Summary", 2)
add_para(
    "The normalised frame-level data are divided into sliding windows (default: 60 frames / "
    "15-frame stride, matching the BehaviorModelConfig). For each window, the following "
    "summary statistics are computed over every feature column:"
)
stats = ["mean", "std", "median", "max", "10th percentile (p10)", "90th percentile (p90)",
         "energy (sum of squares / window length)", "periodicity (peak FFT magnitude of the "
         "mean-centred signal, capturing rhythmic components)"]
for s in stats:
    add_bullet(s)

add_para(
    "This produces a segment feature vector of length 8 × n_frame_features. For a typical "
    "setup with 9 pose features and ~14 context features (~23 total), this yields a ~184-dimensional "
    "segment vector. The full merged representation is cached in derived/representations/ "
    "as two Parquet files: frame_features.parquet and segment_features.parquet. A cache "
    "validity check detects stale segment caches when new sessions are added."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. ACTIVE LEARNING TAB
# ══════════════════════════════════════════════════════════════════════════════
add_heading("13. Active Learning Tab: Closed-Loop Modeling Pipeline", 1)

add_para(
    "The Active Learning tab is the core of ABEL. It orchestrates a multi-stage pipeline "
    "that runs in a background thread, builds or updates a behavior classifier, and generates "
    "a prioritised review queue for the next human-labeling round. Each run produces an updated "
    "model and a refreshed candidate list."
)

# ── 13.1 Pipeline Stages ───────────────────────────────────────────────────
add_heading("13.1 Pipeline Stages", 2)

add_para("Each pipeline run executes the following stages in order:")

stages = [
    ("1. Context Feature Extraction", "Runs ContextFeatureService on any sessions whose context "
     "features are missing or stale. Required for the representation step. Can be skipped if "
     "already computed from the Context Features tab."),
    ("2. Representation Building", "Merges pose + context features and computes segment-level "
     "summary vectors (BehaviorRepresentationService). Uses the cached representation when "
     "valid; rebuilds automatically when cache is stale."),
    ("3. Training Set Assembly", "Loads reviewer_labels.parquet and joins label records to "
     "segment feature rows. Segments labeled 'accept' (or the target behavior name) are "
     "encoded as positive; rejects become 'no_behavior' negatives. The training set is "
     "accumulated across all review rounds and saved as a versioned Parquet snapshot."),
    ("4. Temporal False-Positive Feedback", "Before training, any training segments whose temporal "
     "centre falls within user-flagged false-positive intervals (from the Temporal Review tab) "
     "are relabeled to 'no_behavior'. This closes the loop between dense temporal inference and "
     "the active-learning model."),
    ("5. Model Training", "Fits a gradient-boosted classifier on the assembled training set "
     "(see Section 13.2)."),
    ("6. Probability Calibration", "Post-hoc probability calibration of the raw classifier "
     "scores (see Section 13.3)."),
    ("7. Full-Dataset Inference", "The calibrated model scores every segment in the "
     "representation file, producing a prediction_prob and raw probability vector for each segment."),
    ("8. Uncertainty Scoring", "Computes uncertainty metrics from the prediction probabilities "
     "and segment features (see Section 13.4)."),
    ("9. Candidate Ranking", "Ranks unlabeled segments by the configured query strategy "
     "(uncertainty, prototype similarity, novelty, etc.) and selects the top-k for the "
     "next review round (see Section 13.7)."),
    ("10. Video-Based Fusion (optional)", "For segments above the fusion uncertainty threshold, "
     "re-scores predictions by blending the pose-context model score with a video-crop "
     "embedding similarity score (see Section 13.6)."),
    ("11. Evaluation", "Computes held-out precision, recall, F1, and PR-AUC on the validation "
     "split; generates PR curve, confusion matrix, and UMAP/PCA separation plot."),
    ("12. Phase 1 Adaptive Benchmarking (optional)", "When enabled, runs multi-modality "
     "benchmarking to characterise the behavior's discriminability (see Section 18)."),
]
for name, desc in stages:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

# ── 13.2 Classifier Training ───────────────────────────────────────────────
add_heading("13.2 Classifier Training", 2)

add_para(
    "The primary classifier is a gradient-boosted decision tree ensemble. The default and "
    "recommended backend is XGBoost, with automatic fallback to LightGBM and then to "
    "scikit-learn HistGradientBoostingClassifier if those are unavailable. "
    "RandomForestClassifier is also available as an alternative."
)

add_heading("XGBoost Configuration", 3)
add_bullet("Tree method: hist (histogram-based, fast for tabular data)")
add_bullet("Device: CUDA on Windows when a GPU is present; falls back to CPU on failure.")
add_bullet("Training uses sample weights: no_behavior segments receive an upward weight of 1.5× "
           "(configurable) to reduce overconfidence toward the negative class when negatives "
           "outnumber positives.")
add_bullet("Ambiguous and boundary_error labels are excluded from training.")
add_bullet("Maximum samples per class can be capped to limit training time on large datasets.")

add_heading("Train / Validation Split", 3)
add_para(
    "The training set is split into train and validation partitions. Three split strategies are "
    "available, all implemented via GroupShuffleSplit:"
)
add_bullet("group_shuffle_session (default, 25% held out): sessions are treated as groups; the "
           "held-out set contains entire sessions not seen during training. Recommended for most "
           "datasets to assess cross-session generalisation.")
add_bullet("group_shuffle_subject: subjects are groups; evaluates generalisation across animals.")
add_bullet("leave_one_subject_out: one subject is held out entirely per fold.")

add_para(
    "When too few unique groups exist for a clean split (e.g., a bootstrap run with a single "
    "session), the pipeline falls back to a stratified random row split to avoid a hard failure."
)

add_heading("Co-occurring Behavior Support", 3)
add_para(
    "When allow_co_occurring_behaviors is enabled, labels stored as pipe-separated strings "
    "(e.g., 'grooming|rearing') are expanded into separate single-label rows before training "
    "so each constituent behavior receives an independent positive training example from the "
    "same feature vector."
)

# ── 13.3 Calibration ──────────────────────────────────────────────────────
add_heading("13.3 Probability Calibration", 2)

add_para(
    "Raw gradient-boosted classifier scores are often poorly calibrated — probabilities near "
    "0 or 1 may be overconfident. Post-hoc calibration is applied using "
    "sklearn.calibration.CalibratedClassifierCV with cv='prefit' (calibration on the same "
    "training set, not via cross-validation, to preserve sample efficiency)."
)
add_bullet("Sigmoid (Platt scaling, default): fits a logistic function mapping raw scores to "
           "calibrated probabilities. Suitable for most dataset sizes.")
add_bullet("Isotonic regression: non-parametric monotone mapping. More flexible but requires "
           "more calibration data (~1000+ examples).")
add_bullet("None: raw classifier probabilities are used without calibration.")

add_para(
    "Calibration quality is reported at the end of each training run via a reliability diagram "
    "(fraction of positives vs. mean predicted probability per bin, 10 bins)."
)

# ── 13.4 Uncertainty Scoring ──────────────────────────────────────────────
add_heading("13.4 Uncertainty Scoring", 2)

add_para(
    "Each unlabeled segment is assigned a composite uncertainty score used to prioritise "
    "the next review batch. Four uncertainty components are computed:"
)

unc_tbl = doc.add_table(rows=1, cols=3)
unc_tbl.style = 'Table Grid'
unc_tbl.rows[0].cells[0].text = "Component"
unc_tbl.rows[0].cells[1].text = "Formula"
unc_tbl.rows[0].cells[2].text = "Default weight"
for c in unc_tbl.rows[0].cells:
    for run in c.paragraphs[0].runs:
        run.bold = True

unc_rows = [
    ("Entropy", "H = -Σ p_i · log(p_i) over class probabilities", "0.40"),
    ("Ensemble variance", "Mean variance across resample ensemble predictions "
     "(average of per-class variances across bootstrap resamples)", "0.40"),
    ("Density outlier score", "Mean k-NN distance in feature space (k=10) — "
     "high scores indicate segments far from the training distribution", "0.20"),
    ("Margin (supplementary)", "1 − |p_top1 − p_top2| — distance between the two "
     "highest class probabilities, inverted so near-boundary segments score high", "0.00 (off by default)"),
]
for row_data in unc_rows:
    add_table_row(unc_tbl, row_data)

add_para(
    "The composite raw uncertainty is computed as:"
)
add_code("uncertainty_raw = w_entropy × H + w_ensemble_var × Var + w_density × kNN_dist + w_margin × Margin_unc")

add_para(
    "The raw score is min-max normalised to [0, 1] across all segments in the current run. "
    "All four components and the final uncertainty_score are stored in the segment feature "
    "table for inspection in the review interface."
)

# ── 13.5 Context Features in Active Learning ──────────────────────────────
add_heading("13.5 Context Feature Integration in Active Learning", 2)

add_para(
    "In the Active Learning tab, context features from the video (optical-flow, substrate "
    "dynamics, spatial proximity) are integrated into the classifier feature vector at the "
    "segment representation stage (Section 12). This is distinct from the Extracting Features "
    "tab: the pose-only features (Section 6) are used for initial motif discovery and candidate "
    "generation (Sections 8–9), which requires no video. Context features join the classifier "
    "input only when the active-learning pipeline builds the segment representation."
)

add_para(
    "Operationally: the Active Learning pipeline first checks whether context features "
    "exist for each session. If not, it runs ContextFeatureService in-pipeline (Stage 1 of "
    "the pipeline run). Context features are then merged with pose features per frame and "
    "summarised into segment statistics (Stage 2, BehaviorRepresentationService). The "
    "classifier therefore sees a combined feature vector encoding both what the animal's "
    "kinematics look like and what is happening at the animal's body position in the video."
)

add_para(
    "The relative contribution of context features is not fixed: it is learned by the "
    "gradient-boosted classifier via feature importance during training. Feature importance "
    "metrics (gain-based) are stored in the model card for inspection."
)

# ── 13.6 Fusion Inference ─────────────────────────────────────────────────
add_heading("13.6 Video-Based Fusion Inference (Optional)", 2)

add_para(
    "After the classifier scores all segments, an optional fusion step re-scores segments "
    "whose classifier-assigned prediction probability falls within a configurable uncertainty "
    "band (default: prediction_prob between fusion_threshold − delta and fusion_threshold + delta). "
    "These are the 'borderline' segments where additional video information may break the tie."
)

add_para(
    "Two video-embedding backends are available:"
)
add_bullet("R3D-18 (preferred): a pre-trained 3D ResNet-18 (torchvision) extracts a "
           "512-dimensional spatio-temporal embedding from 16 evenly sampled frames of the "
           "animal-centred crop. The embedding is taken from the output of layer2 (before "
           "global pooling) and mean-pooled over the temporal dimension. R3D-18 weights use "
           "the Kinetics-400 pre-trained checkpoint.")
add_bullet("Handcrafted fallback: when PyTorch is unavailable, a 6-dimensional summary vector "
           "is computed: [mean pixel intensity, std intensity, mean gradient magnitude, "
           "std gradient magnitude, mean frame-to-frame absolute difference, "
           "std frame-to-frame difference] averaged over sampled frames.")

add_para(
    "The fused score is a weighted combination: "
    "score_fused = α × pose_context_prob + β × cosine_similarity(video_embedding, positive_centroid), "
    "where α = 0.7 and β = 0.3 by default. The positive_centroid is derived from "
    "the mean video embedding of accepted (labeled positive) segments."
)

# ── 13.7 Candidate Ranking Strategies ─────────────────────────────────────
add_heading("13.7 Candidate Ranking and Query Strategies", 2)

add_para(
    "The active-learning tab exposes several query strategies for ranking unlabeled segments:"
)

strat_tbl = doc.add_table(rows=1, cols=2)
strat_tbl.style = 'Table Grid'
strat_tbl.rows[0].cells[0].text = "Strategy"
strat_tbl.rows[0].cells[1].text = "Description"
for c in strat_tbl.rows[0].cells:
    for run in c.paragraphs[0].runs:
        run.bold = True

strat_rows = [
    ("Uncertainty (default)", "Presents segments with the highest composite uncertainty score. "
     "Targets the model's decision boundary."),
    ("Prototype", "Selects segments most similar (in feature space) to the mean representation "
     "of confirmed positive examples. Finds more examples like known positives."),
    ("Novelty", "Selects segments most dissimilar from all previously reviewed examples. "
     "Maximises coverage of the feature space."),
    ("Low Probability (absent)", "Selects unlabeled segments with low classifier probability "
     "(< configurable threshold, default 0.25). Identifies potential missed detections."),
    ("Random Low-Prob (absent)", "Randomly samples from segments with low classifier probability. "
     "Reduces selection bias compared to pure low-prob."),
]
for row_data in strat_rows:
    add_table_row(strat_tbl, row_data)

add_para(
    "For 'Uncertainty' mode with the weighted queue scoring option enabled, a seven-component "
    "priority score is computed per segment:"
)
add_code(
    "priority = w_candidate × candidate_score\n"
    "           + w_uncertainty × uncertainty_score\n"
    "           + w_disagreement × expert_disagreement_score\n"
    "           + w_diversity × diversity_score\n"
    "           + w_confound × confound_score\n"
    "           + w_hard_neg × hard_negative_score\n"
    "           + w_exploration × exploration_score"
)
add_para(
    "Default weights: candidate 0.35, uncertainty 0.20, disagreement 0.15, diversity 0.10, "
    "confound 0.10, hard negatives 0.07, exploration 0.03."
)

add_para(
    "Hard negative mining is triggered when ≥ 8 false positives have accumulated. Hard "
    "negatives are segments predicted positive but rejected by the reviewer — they are "
    "identified by their cluster membership and are over-sampled in subsequent candidate "
    "queues to improve the model's specificity."
)

add_para(
    "A configurable exploration fraction (default: 0.15) of the candidate queue is reserved "
    "for randomly sampled unlabeled segments regardless of their model score, ensuring that "
    "the active-learning process does not collapse to an exploitation-only regime."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. REVIEW TAB
# ══════════════════════════════════════════════════════════════════════════════
add_heading("14. Review Tab: Human-in-the-Loop Labeling", 1)

add_para(
    "The Review tab is the primary human–model interaction interface. It presents the "
    "top-ranked candidate clips from the active-learning pipeline and collects reviewer "
    "decisions that feed back into the next training round."
)

add_heading("Clip Presentation", 2)
add_para(
    "Clips are presented in order of their computed priority score (highest uncertainty or "
    "highest candidate probability, depending on mode). For each clip, the interface shows:"
)
add_bullet("Embedded video player (extracted clip) with configurable playback speed.")
add_bullet("Prediction probability and uncertainty score from the current model.")
add_bullet("Session ID and frame range of the candidate.")
add_bullet("Score component breakdown (candidate score, uncertainty, motif score, seed similarity).")

add_heading("Review Decisions", 2)

rev_tbl = doc.add_table(rows=1, cols=2)
rev_tbl.style = 'Table Grid'
rev_tbl.rows[0].cells[0].text = "Decision"
rev_tbl.rows[0].cells[1].text = "Effect"
for c in rev_tbl.rows[0].cells:
    for run in c.paragraphs[0].runs:
        run.bold = True

rev_rows = [
    ("Accept", "Adds the segment as a positive training example for the target behavior."),
    ("Reject", "Adds the segment as a no_behavior negative."),
    ("Relabel", "Assigns the segment to a different behavior label (for co-occurring or "
     "multi-behavior projects)."),
    ("Ambiguous", "Marks the segment as unresolved; excluded from training."),
    ("Skip", "Defers the segment without a decision."),
    ("Bookmark", "Flags for later review without making a training-relevant decision."),
]
for row_data in rev_rows:
    add_table_row(rev_tbl, row_data)

add_heading("Label Persistence and Fingerprinting", 2)
add_para(
    "Every review decision is written to derived/review_labels/reviewer_labels.parquet. "
    "To detect edits at constant row counts (e.g., a relabel that changes a label without "
    "adding a row), the system maintains a SHA-1 content fingerprint over the sorted "
    "(segment_id, review_label, reviewer_id, confidence, notes) columns. A mismatch triggers "
    "a retrain recommendation."
)

add_heading("Dissimilarity-Based Outlier Detection", 2)
add_para(
    "The review tab optionally runs a dissimilarity analysis that flags candidate segments "
    "whose feature-space distance from the labeled positive centroid significantly exceeds "
    "the within-positive pairwise distance distribution. These outliers may represent "
    "mislabeled examples or segment-boundary errors and are surfaced with a visual warning."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 15. TEMPORAL REFINEMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("15. Temporal Refinement: Dense Sliding-Window Inference", 1)

add_para(
    "After a satisfactory classifier has been trained via active learning, the Temporal "
    "Refinement step applies it at high temporal resolution across entire recordings. Unlike "
    "the active-learning pipeline (which uses a fixed 60-frame stride), dense inference uses "
    "overlapping windows stepping by a single configurable interval (default: 0.1 s = 3 frames "
    "at 30 fps) to produce a smooth per-frame probability trace for each behavior."
)

add_heading("Dense Inference Procedure", 2)
add_para(
    "For each recording session and behavior, the service:"
)
add_numbered("Loads the active-learning model (model_state.pkl) and the full frame-level "
             "feature representation.")
add_numbered("Constructs a dense sequence of overlapping windows at the configured step size, "
             "starting after a warm-up period (default: 1.5 s) to allow context features to "
             "stabilise.")
add_numbered("Runs batch inference on all windows using the trained XGBoost / LightGBM / "
             "HistGBDT model. XGBoost inference uses the DMatrix code path to avoid the "
             "'mismatched devices' warning on Windows when the booster was trained on CUDA.")
add_numbered("Averages the predicted positive probability across all windows overlapping each "
             "frame to produce a per-frame mean probability, with the number of contributing "
             "windows also recorded.")
add_numbered("Applies mutual-inhibition smoothing (see below) across all active behaviors.")
add_numbered("Writes per-session probability traces to "
             "derived/temporal_refinement/{behavior_id}/{session_id}_probabilities.parquet.")

add_heading("Mutual Inhibition", 2)
add_para(
    "When multiple behaviors are modeled simultaneously, a subtractive mutual-inhibition step "
    "penalises frames where more than one behavior is simultaneously highly probable. For each "
    "frame and each behavior b:"
)
add_code(
    "inhibition_b[t] = max over all other behaviors b' of prob_b'[t]\n"
    "prob_b_inhibited[t] = prob_b[t] - inhibition_weight × inhibition_b[t]\n"
    "prob_b_inhibited[t] = clip(prob_b_inhibited[t], 0, 1)"
)
add_para(
    "Default inhibition weight: 0.20. This implements a soft winner-take-all dynamic that "
    "reduces spurious simultaneous detections of mutually exclusive behaviors, while the "
    "allow_co_occurring_behaviors flag disables this inhibition for behaviors that may "
    "legitimately co-occur."
)

add_heading("Parallel Worker Dispatch", 2)
add_para(
    "Temporal refinement supports parallel processing across sessions using a ProcessPoolExecutor "
    "(default: auto-detected number of workers = n_CPU_cores − 1). Each worker handles one "
    "session independently, reading its own copy of the model and features. A shared cancellation "
    "event allows in-progress workers to be stopped cleanly."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 16. TEMPORAL REVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading("16. Temporal Review Tab: Bout Extraction and Threshold Tuning", 1)

add_para(
    "The Temporal Review tab provides interactive tools for transforming the continuous "
    "probability trace into a discrete set of behavioral bouts."
)

add_heading("Probability Smoothing", 2)
add_para(
    "A configurable smoothing step is applied to the raw probability trace before thresholding:"
)
add_bullet("Moving average (default): centred rectangular kernel of configurable width (default: 5 frames).")
add_bullet("Gaussian: Gaussian-weighted kernel.")
add_bullet("Savitzky-Golay: polynomial smoothing for better edge preservation.")

add_heading("Bout Extraction", 2)
add_para(
    "The smoothed probability trace is thresholded with hysteresis:"
)
add_bullet("Onset threshold (default: 0.50): probability must rise above this value to begin a bout.")
add_bullet("Offset threshold (default: same as onset): probability must fall below this value to end a bout.")
add_bullet("Minimum bout duration (default: 6 frames ≈ 200 ms at 30 fps): bouts shorter than this "
           "are discarded.")
add_bullet("Merge gap (default: 3 frames): bouts separated by a gap shorter than this are merged.")

add_heading("False-Positive Feedback", 2)
add_para(
    "Detected bouts shown to the researcher in the Temporal Review tab can be flagged as "
    "false positives. Flagged frame intervals are stored in "
    "derived/temporal_refinement/{behavior_id}/feedback_intervals.json. On the next active-"
    "learning training run, any training segments whose temporal centre falls within a "
    "flagged interval are relabeled to no_behavior, directly incorporating the temporal "
    "review feedback into the classifier."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 17. EVALUATION METRICS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("17. Model Evaluation Metrics", 1)

add_para(
    "ABEL reports evaluation metrics at three levels of temporal resolution:"
)

add_heading("Segment-Level Metrics", 2)
add_para(
    "Computed on the held-out validation split from the active-learning train/val split:"
)
add_bullet("Precision: TP / (TP + FP)")
add_bullet("Recall: TP / (TP + FN)")
add_bullet("Macro F1: harmonic mean of precision and recall, averaged across classes.")
add_bullet("PR-AUC (Average Precision): area under the precision-recall curve, computed "
           "for the target behavior class. This is the primary metric reported in the "
           "training log and model card.")
add_bullet("Confusion matrix (visualised as a PNG heatmap).")
add_bullet("Calibration curve: fraction of positives vs. mean predicted probability across "
           "10 bins (reliability diagram).")
add_bullet("Expected calibration error (ECE): weighted mean absolute deviation between "
           "predicted probability and observed positive rate.")

add_heading("Frame-Level Metrics", 2)
add_para(
    "The same precision/recall/F1 and PR-AUC metrics are computed at the individual frame "
    "level using the expanded prediction-probability trace from dense inference."
)

add_heading("Bout-Level Metrics", 2)
add_para(
    "After bout extraction, the following behavioral assay metrics are computed from the merged "
    "bout table:"
)
add_bullet("Latency to first bout (seconds)")
add_bullet("Total behavior time (seconds)")
add_bullet("Number of bouts")
add_bullet("Mean bout duration (seconds)")
add_bullet("Behavior rate (bouts per second of recording)")
add_bullet("Duration-weighted mean distance to target ROI during bouts (mm, when calibrated)")

add_heading("Visualization Outputs", 2)
add_bullet("PR_curve.png: precision-recall curve.")
add_bullet("confusion_matrix.png: annotated confusion matrix heatmap.")
add_bullet("behavior_separation_active_learning.png: UMAP (or PCA fallback) scatter plot "
           "of the validation-set segment features, coloured by class label.")

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 18. PHASE 1 ADAPTIVE BENCHMARKING
# ══════════════════════════════════════════════════════════════════════════════
add_heading("18. Phase 1 Adaptive Benchmarking (Advanced)", 1)

add_para(
    "Phase 1 benchmarking is an opt-in diagnostic step that characterises how well the "
    "behavior can be distinguished from non-behavior across different feature modalities and "
    "temporal scales before the full active-learning loop converges. It is primarily intended "
    "for publication-quality diagnostics and for troubleshooting datasets where the classifier "
    "achieves low performance."
)

add_heading("Modality Benchmarking", 2)
add_para(
    "The service trains independent XGBoost classifiers on separate feature families:"
)
add_bullet("Pose-only features")
add_bullet("Context-only features (optical flow, substrate motion)")
add_bullet("Spatial features (distance/angle to ROI)")
add_bullet("Full combined feature vector")

add_para(
    "Performance (Average Precision at the representative segment window, 0.5 s) is "
    "compared across modalities to identify which feature groups are most informative for "
    "the specific behavior."
)

add_heading("Confound Analysis", 2)
add_para(
    "When sufficient non-target labels are available, a confound classifier is trained "
    "on the non-target label set vs. no_behavior to identify features that distinguish "
    "confounding behaviors from true negatives. Features with high importance for both the "
    "target and confound classifiers are flagged as potential confounds."
)

add_heading("Outputs", 2)
add_bullet("derived/analysis/benchmarks/{behavior_id}/: CSV tables of metrics per modality.")
add_bullet("derived/analysis/diagnostics/{behavior_id}/: publication-quality PNG and SVG plot "
           "files showing modality comparison bar charts and calibration "
           "reliability diagrams.")

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 19. EXPORT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("19. Export", 1)

add_para(
    "The Export tab provides three output formats:"
)

add_heading("CSV Export (per-session)", 2)
add_para(
    "One row per confirmed bout, columns: subject_id, session_id, behavior, start_frame, "
    "end_frame, start_time_s, end_time_s, duration_s, confidence."
)

add_heading("XLSX Workbook Export (per-subject sheets)", 2)
add_para(
    "Each subject gets a separate Excel worksheet. Rows represent time segments; columns "
    "represent behaviors. Binary (0/1) or probability values are written per frame or per "
    "bout. Subject ordering follows the user-defined order set in the export settings."
)

add_heading("Bout Frame Export", 2)
add_para(
    "Start and end frame numbers for each accepted bout are exported, suitable for "
    "downstream ethogram alignment or frame-accurate comparison with other scorers."
)

add_para(
    "All behavior names are resolved via the BehaviorService alias map (behavior_id → "
    "display name → short name) so that outputs use human-readable labels regardless of "
    "the internal ID used during training."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 20. PROVENANCE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("20. Provenance and Reproducibility", 1)

add_para(
    "Every model artifact, evaluation report, and derived feature file includes an "
    "ArtifactProvenance record containing:"
)
add_bullet("app_version: ABEL semantic version.")
add_bullet("git_commit_hash: repository commit at time of artifact generation.")
add_bullet("model_version: the versioned model identifier (e.g., 'behavior_model_v1').")
add_bullet("feature_version: the versioned feature representation identifier.")
add_bullet("config_hash: SHA-256 hash of the project configuration YAML at time of run.")
add_bullet("timestamp: UTC datetime of artifact creation.")

add_para(
    "Training set snapshots are stored with timestamps in "
    "derived/training_sets/snapshots/, so the exact training data for any model version "
    "can be reconstructed. WorkflowSnapshot records capture the complete pipeline "
    "configuration at each run for audit purposes."
)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 21. DEPENDENCIES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("21. Software Dependencies", 1)

add_para(
    "ABEL uses a tiered dependency model. Core dependencies are required for the GUI "
    "to launch; heavy ML and computer vision packages are installed on demand."
)

dep_tbl = doc.add_table(rows=1, cols=4)
dep_tbl.style = 'Table Grid'
dep_tbl.rows[0].cells[0].text = "Package"
dep_tbl.rows[0].cells[1].text = "Tier"
dep_tbl.rows[0].cells[2].text = "Purpose"
dep_tbl.rows[0].cells[3].text = "Fallback"
for c in dep_tbl.rows[0].cells:
    for run in c.paragraphs[0].runs:
        run.bold = True

dep_rows = [
    ("PySide6 ≥ 6.7", "Core", "Qt6 GUI framework", "Required"),
    ("pydantic ≥ 2.7", "Core", "Data validation and serialisation", "Required"),
    ("numpy ≥ 1.26", "Core", "Numerical arrays", "Required"),
    ("pandas ≥ 2.2", "Core", "Tabular data and Parquet I/O", "Required"),
    ("PyYAML ≥ 6.0", "Core", "YAML config files", "Required"),
    ("scikit-learn ≥ 1.4", "Tier 2", "Clustering, calibration, metrics", "HistGBDT fallback"),
    ("xgboost ≥ 2.0", "Tier 2", "Primary gradient-boosted classifier (GPU-capable)", "LightGBM → HistGBDT"),
    ("lightgbm ≥ 4.0", "Tier 2", "Alternative GBM classifier", "HistGBDT"),
    ("opencv-python ≥ 4.9", "Tier 2", "Video decoding, optical flow, background subtraction", "Context features disabled"),
    ("torch ≥ 2.2", "Tier 3", "R3D-18 video embedding for fusion inference", "Handcrafted fallback"),
    ("torchvision ≥ 0.17", "Tier 3", "R3D-18 model weights", "Handcrafted fallback"),
    ("umap-learn ≥ 0.5", "Tier 3", "UMAP dimensionality reduction for motif discovery and evaluation", "PCA fallback"),
    ("hdbscan ≥ 0.8", "Tier 3", "Density-based clustering for motif discovery", "K-Means fallback"),
    ("matplotlib ≥ 3.8", "Tier 3", "PR curve, confusion matrix, UMAP/PCA plots", "Plots disabled"),
    ("pyarrow ≥ 15.0", "Tier 2", "Parquet serialisation for feature files", "Required for Parquet"),
    ("h5py ≥ 3.10", "Tier 2", "Reading DLC H5 pose files", "CSV-only mode"),
]
for row_data in dep_rows:
    add_table_row(dep_tbl, row_data)

section_break()

# ══════════════════════════════════════════════════════════════════════════════
# 22. GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("22. Glossary", 1)

glossary = [
    ("Active learning", "An iterative machine-learning paradigm in which the model queries "
     "the most informative unlabeled examples for human annotation, reducing the total "
     "labeling burden relative to random sampling."),
    ("Bout", "A contiguous temporal interval during which a behavior is classified as 'on'. "
     "Bouts are extracted from the probability trace by thresholding with optional hysteresis, "
     "minimum duration, and gap-merge parameters."),
    ("Body centroid", "The mean position of all tracked keypoints with above-threshold likelihood "
     "at a given frame. Used as the crop centre for clip extraction and as the trajectory "
     "reference for kinematic features."),
    ("Candidate segment / window", "A short temporal interval (default: 2 s) selected by the "
     "pipeline for human review based on its ranked priority score."),
    ("Calibration (probability)", "Post-hoc adjustment of classifier output scores to align "
     "predicted probabilities with empirical positive rates."),
    ("Context features", "Per-frame features derived from the raw video (optical flow, substrate "
     "dynamics, spatial proximity to ROI) as opposed to pose-tracking kinematics."),
    ("DLC / DeepLabCut", "An open-source markerless pose estimation framework that tracks user-"
     "defined body keypoints in video, outputting x/y coordinates and likelihood per frame."),
    ("Ensemble variance", "The disagreement across multiple predictions of the same example "
     "(via bootstrap resampling) used as a model uncertainty measure."),
    ("Episode", "A variable-length candidate segment constructed by merging consecutive "
     "same-motif windows from the sliding-window feature extraction, prior to clip extraction."),
    ("Feature representation", "The tabular feature matrix (segment_features.parquet) formed "
     "by combining per-group-normalised pose and context features, summarised over sliding "
     "windows of configurable length and stride."),
    ("Farnebäck optical flow", "A dense optical flow algorithm that fits a polynomial expansion "
     "to the image intensities to estimate per-pixel displacement between frames."),
    ("Hard negatives", "Segments predicted positive by the model but labeled negative by the "
     "reviewer. Deliberately over-sampled in subsequent review queues to improve specificity."),
    ("Kinematic feature", "A scalar descriptor of the animal's movement trajectory computed "
     "from pose tracking: speed, displacement, body-axis angle, angular variability."),
    ("Likelihood threshold", "A confidence floor applied to DLC keypoint detections. Frames "
     "below this value are treated as missing and interpolated."),
    ("MOG2", "Mixture of Gaussians background subtractor (OpenCV). Maintains a running model "
     "of the background appearance and outputs a foreground mask."),
    ("Motif", "A recurrent kinematic pattern identified by unsupervised clustering of pose "
     "feature windows. Motifs are not labeled behaviors; they serve as a data-driven prior "
     "for candidate generation."),
    ("Mutual inhibition", "A soft winner-take-all post-processing step applied to multi-behavior "
     "probability traces to reduce spurious simultaneous detections."),
    ("PR-AUC (Average Precision)", "The area under the precision-recall curve, summarising "
     "a classifier's performance across all decision thresholds. Preferred over ROC-AUC for "
     "imbalanced datasets."),
    ("Provenance", "The complete record of software version, configuration hash, and timestamps "
     "associated with a derived artifact, enabling future reproduction of results."),
    ("R3D-18", "A 3D convolutional ResNet-18 pre-trained on Kinetics-400 used in ABEL as "
     "a spatio-temporal video embedding model for fusion inference."),
    ("ROI (Region of Interest)", "A user-defined rectangular region in the video frame "
     "(e.g., the location of a stimulus object) used to compute spatial context features "
     "such as distance-to-target."),
    ("Seed example", "A user-annotated positive instance of the target behavior providing "
     "the initial supervisory signal before any rated active-learning cycles."),
    ("Session", "A single video + pose file pair corresponding to one recording of one animal."),
    ("Uncertainty sampling", "An active-learning query strategy that prioritises examples "
     "near the classifier's decision boundary (highest entropy / lowest margin)."),
    ("UMAP", "Uniform Manifold Approximation and Projection — a non-linear dimensionality "
     "reduction algorithm used in ABEL for motif discovery and evaluation visualisation."),
    ("Validation split", "The held-out portion of the labeled training set (default: 25%) "
     "used exclusively for performance evaluation and not for model fitting or calibration."),
    ("XGBoost", "eXtreme Gradient Boosting — an efficient open-source gradient-boosted "
     "decision tree library with CPU and GPU (CUDA) training support."),
    ("Z-score normalisation", "Standardisation of feature values to zero mean and unit "
     "variance within each (animal, session) group, removing systematic scale differences "
     "across recording conditions."),
]

for term, definition in sorted(glossary, key=lambda x: x[0].lower()):
    p = doc.add_paragraph()
    run = p.add_run(f"{term}: ")
    run.bold = True
    p.add_run(definition)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = r"c:\Users\jober\Desktop\ABEL realism\ABEL_Methods_Documentation.docx"
doc.save(out_path)
print(f"Document saved to: {out_path}")
